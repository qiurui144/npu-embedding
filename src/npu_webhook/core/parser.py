"""文件解析器：Markdown / 纯文本 / 代码 / PDF / DOCX"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# 代码文件扩展名
CODE_EXTENSIONS = {
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".go", ".rs",
    ".c", ".cpp", ".h", ".hpp", ".cs", ".rb", ".php", ".sh",
    ".yaml", ".yml", ".toml", ".json", ".xml", ".sql",
}


def parse_file(path: str | Path) -> tuple[str, str]:
    """解析文件内容，返回 (title, content)

    支持: .md, .txt, 代码文件, .pdf, .docx
    """
    path = Path(path)
    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            return _parse_pdf(path)
        elif suffix == ".docx":
            return _parse_docx(path)
        elif suffix == ".md":
            return _parse_markdown(path)
        elif suffix in CODE_EXTENSIONS:
            return _parse_code(path)
        else:
            return _parse_text(path)
    except Exception:
        logger.exception("Failed to parse file: %s", path)
        return path.name, ""


def _parse_markdown(path: Path) -> tuple[str, str]:
    """解析 Markdown，提取第一个标题作为 title"""
    content = path.read_text(encoding="utf-8", errors="replace")
    title = path.stem
    for line in content.splitlines():
        line = line.strip()
        if line.startswith("# "):
            title = line[2:].strip()
            break
    return title, content


def _parse_text(path: Path) -> tuple[str, str]:
    """解析纯文本"""
    content = path.read_text(encoding="utf-8", errors="replace")
    title = path.stem
    first_line = content.strip().split("\n", 1)[0][:100] if content.strip() else path.stem
    if first_line:
        title = first_line
    return title, content


def _parse_code(path: Path) -> tuple[str, str]:
    """解析代码文件，保留原始内容"""
    content = path.read_text(encoding="utf-8", errors="replace")
    return f"{path.name}", content


def _parse_pdf(path: Path) -> tuple[str, str]:
    """解析 PDF（使用 PyMuPDF）"""
    import pymupdf

    doc = pymupdf.open(str(path))
    title = doc.metadata.get("title", "") or path.stem
    pages = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text)
    doc.close()
    return title, "\n\n".join(pages)


def _parse_docx(path: Path) -> tuple[str, str]:
    """解析 DOCX"""
    from docx import Document

    doc = Document(str(path))
    title = path.stem
    # 尝试从核心属性获取标题
    if doc.core_properties.title:
        title = doc.core_properties.title
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return title, "\n\n".join(paragraphs)


def parse_bytes(data: bytes, filename: str) -> tuple[str, str]:
    """从内存 bytes 解析文件，返回 (title, content)。

    filename 仅用于类型检测（扩展名），不做磁盘操作。
    复用现有各格式解析逻辑。
    """
    suffix = Path(filename).suffix.lower()
    name_stem = Path(filename).stem

    try:
        if suffix == ".pdf":
            return _parse_pdf_bytes(data, name_stem)
        elif suffix == ".docx":
            return _parse_docx_bytes(data, name_stem)
        elif suffix == ".md":
            content = data.decode("utf-8", errors="replace")
            title = name_stem
            for line in content.splitlines():
                if line.strip().startswith("# "):
                    title = line.strip()[2:].strip()
                    break
            return title, content
        elif suffix in CODE_EXTENSIONS:
            return filename, data.decode("utf-8", errors="replace")
        else:
            content = data.decode("utf-8", errors="replace")
            title = content.strip().split("\n", 1)[0][:100] if content.strip() else name_stem
            return title or name_stem, content
    except Exception:
        logger.exception("Failed to parse bytes for: %s", filename)
        return name_stem, ""


def _parse_pdf_bytes(data: bytes, name_stem: str) -> tuple[str, str]:
    """从内存 bytes 解析 PDF（使用 PyMuPDF）"""
    import io

    import pymupdf

    doc = pymupdf.open(stream=io.BytesIO(data), filetype="pdf")
    title = doc.metadata.get("title", "") or name_stem
    pages = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text)
    doc.close()
    return title, "\n\n".join(pages)


def _parse_docx_bytes(data: bytes, name_stem: str) -> tuple[str, str]:
    """从内存 bytes 解析 DOCX"""
    import io

    from docx import Document

    doc = Document(io.BytesIO(data))
    title = doc.core_properties.title or name_stem
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return title, "\n\n".join(paragraphs)
